from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Sequence

from docx import Document


@dataclass
class LoadedText:
    """Контейнер для загруженного текста."""

    path: str
    paragraphs: List[str]

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)


def load_docx(path: str | Path) -> LoadedText:
    """Читает .docx и возвращает непустые параграфы."""
    p = Path(path)
    doc = Document(str(p))

    paragraphs: List[str] = []
    for par in doc.paragraphs:
        t = par.text
        if t and t.strip():
            paragraphs.append(t)

    return LoadedText(path=str(p), paragraphs=paragraphs)


def load_many_docx(paths: Sequence[str | Path]) -> List[LoadedText]:
    return [load_docx(p) for p in paths]
